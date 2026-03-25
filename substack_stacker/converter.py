import io
import re
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from substack_stacker.client import download_image


class HtmlToDocxConverter:
    def __init__(self, doc, session):
        self.doc = doc
        self.session = session
        self._current_paragraph = None
        self._style_stack = []  # list of dicts: {"bold": True}, {"italic": True}, etc.
        self._list_depth = 0

    def convert(self, html):
        if not html:
            return
        soup = BeautifulSoup(html, "html.parser")
        self._process_children(soup)
        self._current_paragraph = None

    def _process_children(self, element):
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text.strip() or text == " ":
                    self._add_run(text)
            elif isinstance(child, Tag):
                self._process_tag(child)

    def _process_tag(self, tag):
        name = tag.name.lower()

        # Headings
        if re.match(r"^h([1-6])$", name):
            level = int(name[1])
            self._flush_paragraph()
            heading = self.doc.add_heading(level=level)
            self._current_paragraph = heading
            self._process_children(tag)
            self._flush_paragraph()
            return

        # Paragraph
        if name == "p":
            self._flush_paragraph()
            self._current_paragraph = self.doc.add_paragraph()
            self._process_children(tag)
            self._flush_paragraph()
            return

        # Bold
        if name in ("strong", "b"):
            self._style_stack.append("bold")
            self._process_children(tag)
            self._style_stack.pop()
            return

        # Italic
        if name in ("em", "i"):
            self._style_stack.append("italic")
            self._process_children(tag)
            self._style_stack.pop()
            return

        # Inline code
        if name == "code" and (not tag.parent or tag.parent.name != "pre"):
            self._style_stack.append("code")
            self._process_children(tag)
            self._style_stack.pop()
            return

        # Links
        if name == "a":
            href = tag.get("href", "")
            text = tag.get_text()
            self._ensure_paragraph()
            run = self._current_paragraph.add_run(text)
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.font.underline = True
            self._apply_styles(run)
            # Append URL if it differs from the display text
            if href and href != text and not href.startswith("javascript:"):
                url_run = self._current_paragraph.add_run(f" ({href})")
                url_run.font.size = Pt(8)
                url_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            return

        # Images
        if name == "img":
            src = tag.get("src", "")
            alt = tag.get("alt", "")
            if src:
                self._flush_paragraph()
                data = download_image(self.session, src)
                if data:
                    try:
                        self.doc.add_picture(io.BytesIO(data), width=Inches(5.5))
                    except Exception:
                        self.doc.add_paragraph(f"[Image: {alt or src}]")
                else:
                    self.doc.add_paragraph(f"[Image: {alt or src}]")
                # Add caption if alt text exists
                if alt:
                    caption = self.doc.add_paragraph(alt)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = caption.runs[0] if caption.runs else caption.add_run(alt)
                    run.font.size = Pt(9)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            return

        # Blockquote
        if name == "blockquote":
            self._flush_paragraph()
            # Process children but with blockquote styling
            saved_para = self._current_paragraph
            for child in tag.children:
                if isinstance(child, Tag) and child.name == "p":
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    self._current_paragraph = p
                    self._style_stack.append("italic")
                    self._process_children(child)
                    self._style_stack.pop()
                    self._flush_paragraph()
                elif isinstance(child, Tag):
                    self._process_tag(child)
                elif isinstance(child, NavigableString) and str(child).strip():
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    self._current_paragraph = p
                    self._style_stack.append("italic")
                    self._add_run(str(child))
                    self._style_stack.pop()
                    self._flush_paragraph()
            self._current_paragraph = saved_para
            return

        # Unordered list
        if name == "ul":
            self._flush_paragraph()
            self._process_list(tag, ordered=False)
            return

        # Ordered list
        if name == "ol":
            self._flush_paragraph()
            self._process_list(tag, ordered=True)
            return

        # List item (handled by _process_list, but in case it appears alone)
        if name == "li":
            self._flush_paragraph()
            style = "List Bullet" if self._list_depth == 0 else "List Bullet 2"
            self._current_paragraph = self.doc.add_paragraph(style=style)
            self._process_children(tag)
            self._flush_paragraph()
            return

        # Preformatted / code block
        if name == "pre":
            self._flush_paragraph()
            # Extract text content, preserving whitespace
            code_tag = tag.find("code")
            text = code_tag.get_text() if code_tag else tag.get_text()
            p = self.doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            # Add light gray background via XML
            shd = run._element.get_or_add_rPr()
            shd_elem = shd.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "F0F0F0",
            })
            shd.append(shd_elem)
            self._current_paragraph = None
            return

        # Line break
        if name == "br":
            self._ensure_paragraph()
            run = self._current_paragraph.add_run()
            run.add_break()
            return

        # Horizontal rule
        if name == "hr":
            self._flush_paragraph()
            p = self.doc.add_paragraph("───────────────────────────────")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            self._current_paragraph = None
            return

        # Figure
        if name == "figure":
            self._flush_paragraph()
            self._process_children(tag)
            return

        # Figcaption
        if name == "figcaption":
            self._flush_paragraph()
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._current_paragraph = p
            self._style_stack.append("italic")
            self._process_children(tag)
            self._style_stack.pop()
            if self._current_paragraph and self._current_paragraph.runs:
                for run in self._current_paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            self._flush_paragraph()
            return

        # Subscript / superscript
        if name == "sub":
            self._ensure_paragraph()
            run = self._current_paragraph.add_run(tag.get_text())
            run.font.subscript = True
            return
        if name == "sup":
            self._ensure_paragraph()
            run = self._current_paragraph.add_run(tag.get_text())
            run.font.superscript = True
            return

        # Transparent wrappers: div, section, article, span, main, etc.
        self._process_children(tag)

    def _process_list(self, list_tag, ordered=False):
        self._list_depth += 1
        for child in list_tag.children:
            if not isinstance(child, Tag):
                continue
            if child.name == "li":
                # Check for nested lists
                nested_list = child.find(["ul", "ol"], recursive=False)
                if self._list_depth <= 1:
                    style = "List Number" if ordered else "List Bullet"
                else:
                    style = "List Number 2" if ordered else "List Bullet 2"
                self._current_paragraph = self.doc.add_paragraph(style=style)
                # Process inline content (skip nested lists for now)
                for li_child in child.children:
                    if isinstance(li_child, Tag) and li_child.name in ("ul", "ol"):
                        continue  # handled below
                    elif isinstance(li_child, Tag):
                        self._process_tag(li_child)
                    elif isinstance(li_child, NavigableString):
                        text = str(li_child)
                        if text.strip() or text == " ":
                            self._add_run(text)
                self._flush_paragraph()
                # Now process nested lists
                if nested_list:
                    is_ordered = nested_list.name == "ol"
                    self._process_list(nested_list, ordered=is_ordered)
        self._list_depth -= 1

    def _ensure_paragraph(self):
        if self._current_paragraph is None:
            self._current_paragraph = self.doc.add_paragraph()

    def _flush_paragraph(self):
        self._current_paragraph = None

    def _add_run(self, text):
        self._ensure_paragraph()
        run = self._current_paragraph.add_run(text)
        self._apply_styles(run)
        return run

    def _apply_styles(self, run):
        for style in self._style_stack:
            if style == "bold":
                run.bold = True
            elif style == "italic":
                run.italic = True
            elif style == "code":
                run.font.name = "Courier New"
                run.font.size = Pt(9)
