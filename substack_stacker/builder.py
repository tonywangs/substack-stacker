import time
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tqdm import tqdm

from substack_stacker.client import _make_session, fetch_post_body
from substack_stacker.converter import HtmlToDocxConverter


def build_document(subdomain, posts, output_path, delay=1.5):
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Georgia"
    font.size = Pt(11)

    _add_title_page(doc, subdomain)
    _add_toc(doc, posts)

    session = _make_session()
    converter = HtmlToDocxConverter(doc, session)

    for i, post in enumerate(tqdm(posts, desc="Building document", unit=" posts")):
        doc.add_page_break()

        # Post title
        title = post.get("title", "Untitled")
        heading = doc.add_heading(title, level=1)

        # Subtitle
        subtitle = post.get("subtitle", "")
        if subtitle:
            sub = doc.add_heading(subtitle, level=2)

        # Metadata line: date + author
        date_str = _format_date(post.get("post_date", ""))
        author_name = ""
        if post.get("publishedBylines"):
            author_name = post["publishedBylines"][0].get("name", "")

        meta_parts = []
        if author_name:
            meta_parts.append(author_name)
        if date_str:
            meta_parts.append(date_str)
        if meta_parts:
            meta_para = doc.add_paragraph()
            run = meta_para.add_run(" · ".join(meta_parts))
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # Add spacing before body
        doc.add_paragraph()

        # Fetch and convert body
        slug = post.get("slug", "")
        if slug:
            body_html = fetch_post_body(session, subdomain, slug)
            if body_html:
                converter.convert(body_html)

        if i < len(posts) - 1:
            time.sleep(delay)

    doc.save(output_path)
    print(f"\nSaved {len(posts)} posts to {output_path}")


def _add_title_page(doc, subdomain):
    # Vertical spacing
    for _ in range(6):
        doc.add_paragraph()

    # Publication name
    title = doc.add_heading(subdomain.replace("-", " ").title(), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Compiled from Substack")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()


def _add_toc(doc, posts):
    doc.add_heading("Table of Contents", level=1)

    for i, post in enumerate(posts, 1):
        title = post.get("title", "Untitled")
        date_str = _format_date(post.get("post_date", ""))

        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {title}")
        run.bold = True
        if date_str:
            date_run = p.add_run(f"  —  {date_str}")
            date_run.font.size = Pt(9)
            date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()


def _format_date(date_str):
    if not date_str:
        return ""
    try:
        dt = datetime.fromisoformat(date_str.replace("Z", "+00:00"))
        return dt.strftime("%B %d, %Y")
    except (ValueError, AttributeError):
        return date_str
